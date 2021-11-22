#Imports
from selenium import webdriver
import time
from datetime import datetime
import os
from docx import Document
from hyperlink import add_hyperlink


#Initializations
wait_time = 5
main_url = 'https://punchng.com/'

#open chrome
op = webdriver.ChromeOptions()
op.add_argument('headless')
driver = "chromedriver.exe"
driver = webdriver.Chrome(driver,options=op)
driver.implicitly_wait(wait_time)
driver.get(main_url)
time.sleep(wait_time)

#Instatiate document
document = Document()

#Header
header = driver.find_element_by_xpath('.//header[@id="masthead"]')
header = header.find_element_by_xpath('.//h1[@class="site-title"]')
header = header.find_element_by_xpath('.//a')
header = header.get_attribute('innerHTML')

#Date
date = datetime.today().strftime('%Y-%m-%d')

document.add_heading(header+'\n'+date, level=1)

#Main page
group0 = driver.find_elements_by_xpath('.//div[@class="group-section group-section-splits"]')
group4 = driver.find_element_by_xpath('.//div[@class="row row-home-intro group-type-five columns-grid"]')
group1 = group0[0]
group2 = group0[1]
group3 = group0[2]

# Articles
for gr in [group1,group2,group3,group4]:
    if gr != group4:
        grous = [gr.find_element_by_xpath('.//div[@class="columns small-12 medium-6 large-5"]'),\
        gr.find_element_by_xpath('.//div[@class="columns small-12 medium-6 large-7"]')]
    else:
        grous = group4.find_elements_by_xpath('.//div[@class="columns"]')
    for grou in grous:
        g1=grou.find_element_by_xpath('.//h2[@class="heading-title"]')
        a1 = g1.find_element_by_xpath('.//a')
        title = a1.get_attribute('innerHTML')

        document.add_heading(title, level=2)

        g2=grou.find_elements_by_xpath('.//h3[@class="entry-title"]')
        for g in g2:
            a2 = g.find_element_by_xpath('.//a')
            link = a2.get_attribute("href")
            summary = a2.get_attribute('innerHTML')

            p = document.add_paragraph(style='List Number')

            add_hyperlink(p, link, summary, None, True)

#Save document
path = "../data/punch-news"+date
pathdocx = path+".docx"
pathpdf = path+".pdf"
if os.path.exists(pathdocx):
  os.remove(pathdocx)
else:
  pass
document.save(pathdocx)