# Imports
from selenium import webdriver
import time
from datetime import datetime
import os
from docx import Document
from hyperlink import add_hyperlink

# Initializations
wait_time = 5
main_url = 'https://tvn24.pl/tvn24-news-in-english/'

# open chrome
op = webdriver.ChromeOptions()
op.add_argument('headless')
driver = "chromedriver.exe"
driver = webdriver.Chrome(driver, options=op)
driver.implicitly_wait(wait_time)
driver.get(main_url)
time.sleep(wait_time)

# Instatiate document
document = Document()

# Main page
main_content = driver.find_element_by_xpath('.//section[@class="tvn24-news-in-english main-content-holder"]')

# Header
header = main_content.find_element_by_xpath('.//h1[@class="heading heading--size-30 decorated-header__headline"]').text

# Date
date = datetime.today().strftime('%Y-%m-%d')

document.add_heading(header + '\n' + date, level=1)

# Articles
articles = main_content.find_elements_by_xpath('.//div[@class="teaser-wrapper"]')

for i in range(len(articles)):
    # Ttle
    title = articles[i].find_element_by_xpath('.//header[@class="article-header"]').text
    # Link
    link = articles[i].find_element_by_xpath('.//a[@class="default-teaser__link"]').get_attribute("href")
    # Summary
    summary = articles[i].find_element_by_xpath('.//p[@class="article-lead__text"]').text

    t = document.add_heading(level=2)
    document.add_paragraph(summary)
    add_hyperlink(t, link, title, None, True)

# Save document
path = "../data/tvn24-news" + date
pathdocx = path + ".docx"
pathpdf = path + ".pdf"
if os.path.exists(pathdocx):
    os.remove(pathdocx)
else:
    pass
document.save(pathdocx)
